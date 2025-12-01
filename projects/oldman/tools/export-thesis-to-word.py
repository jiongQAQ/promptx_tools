#!/usr/bin/env python3
"""
论文导出工具 - 将JSON格式的论文内容导出为Word文档
按照玉溪师范学院本科生毕业论文格式要求

使用方法:
    python3 export-thesis-to-word.py [项目路径]

示例:
    python3 export-thesis-to-word.py
    python3 export-thesis-to-word.py /path/to/project
"""

import json
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


class ThesisFormatter:
    """论文格式化器 - 应用玉溪师范学院格式规范"""

    def __init__(self, doc):
        self.doc = doc
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        """设置页面格式"""
        section = self.doc.sections[0]

        # 纸张大小: A4
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

        # 页边距
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

        # 页眉页脚距离
        section.header_distance = Cm(1.75)
        section.footer_distance = Cm(1.0)

    def _setup_styles(self):
        """设置样式"""
        styles = self.doc.styles

        # 正文样式
        if 'Normal' in styles:
            normal = styles['Normal']
            normal.font.name = 'Times New Roman'
            normal.font.size = Pt(12)  # 小四
            normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            normal.paragraph_format.line_spacing = 1.25
            normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            normal.paragraph_format.first_line_indent = Pt(0)  # 无首行缩进

        # 一级标题样式
        if 'Heading 1' not in styles:
            heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        else:
            heading1 = styles['Heading 1']

        heading1.font.name = 'Times New Roman'
        heading1.font.size = Pt(16)  # 三号
        heading1.font.bold = True
        heading1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading1.paragraph_format.space_before = Pt(12)  # 段前1行
        heading1.paragraph_format.space_after = Pt(6)    # 段后0.5行
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 二级标题样式
        if 'Heading 2' not in styles:
            heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        else:
            heading2 = styles['Heading 2']

        heading2.font.name = 'Times New Roman'
        heading2.font.size = Pt(15)  # 小三
        heading2.font.bold = True
        heading2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading2.paragraph_format.space_before = Pt(12)
        heading2.paragraph_format.space_after = Pt(12)
        heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_cover(self, info):
        """添加封面"""
        # 学科分类号
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"学科分类号  {info.get('category', 'S20·4060')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 学校名称（书法字体）
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('玉溪师范学院')
        run.font.name = '华文行楷'
        run.font.size = Pt(36)
        run.font.bold = True

        # 空行
        self.doc.add_paragraph()

        # 本科生毕业论文
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('本科生毕业论文')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(22)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 空行
        for _ in range(3):
            self.doc.add_paragraph()

        # 论文信息
        info_items = [
            ('题目', info.get('title', '')),
            ('姓名', info.get('author', '')),
            ('学号', info.get('student_id', '')),
            ('学院', info.get('college', '数学与信息技术学院')),
            ('专业', info.get('major', '')),
            ('导师', info.get('supervisor', '')),
            ('职称', info.get('title_level', ''))
        ]

        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"{label}  {value}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加分页
        self.doc.add_page_break()

    def add_declarations(self):
        """添加声明页"""
        # 原创性声明
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('原创性声明')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        declaration_text = """本人郑重声明：所呈交的学位论文，是本人在导师的指导下，独立进行研究所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已发表或撰写过的科研成果。对本文的研究作出重要贡献的个人和集体，均已在文中以明确方式标明。本声明的法律责任由本人承担。"""

        p = self.doc.add_paragraph(declaration_text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(0)

        # 签名栏
        p = self.doc.add_paragraph()
        run = p.add_run('论文作者签名：________  日期：________')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 使用授权声明
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('关于学位论文使用授权的声明')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        authorization_text = """本人完全了解玉溪师范学院有关保留、使用学位论文的规定，同意学校保留或向国家有关部门或机构送交论文的复印件和电子版，允许论文被查阅和借阅；本人授权玉溪师范学院可以将本学位论文的全部或部分内容编入有关数据库进行检索，可以采用影印、缩印或其他复制手段保存论文和汇编本学位论文。"""

        p = self.doc.add_paragraph(authorization_text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = self.doc.add_paragraph('（保密论文在解密后应遵守此规定）')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p = self.doc.add_paragraph()
        run = p.add_run('论文作者签名：________  导师签名：________  日期：________')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        self.doc.add_page_break()

    def add_toc(self, chapters):
        """添加目录"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('目  录')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)  # 三号
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # TODO: 自动生成目录（需要更新页码）
        # 这里简化处理，实际应该使用Word的TOC域

        self.doc.add_page_break()

    def add_abstract_cn(self, abstract, keywords):
        """添加中文摘要"""
        # 标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('摘  要')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 空一行
        self.doc.add_paragraph()

        # 摘要内容
        p = self.doc.add_paragraph(abstract)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.25

        # 空一行
        self.doc.add_paragraph()

        # 关键词
        p = self.doc.add_paragraph()
        run = p.add_run('关键词：')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        run = p.add_run(keywords)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        self.doc.add_page_break()

    def add_abstract_en(self, abstract, keywords):
        """添加英文摘要"""
        # 标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Abstract')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True

        # 空一行
        self.doc.add_paragraph()

        # 摘要内容
        p = self.doc.add_paragraph(abstract)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 空一行
        self.doc.add_paragraph()

        # 关键词
        p = self.doc.add_paragraph()
        run = p.add_run('Key words: ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True

        run = p.add_run(keywords)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        self.doc.add_page_break()

    def add_chapter(self, chapter_data, level=1):
        """添加章节内容"""
        # 添加标题
        if level == 1:
            heading_style = 'Heading 1'
        elif level == 2:
            heading_style = 'Heading 2'
        else:
            heading_style = 'Normal'

        title = f"{chapter_data.get('id', '')} {chapter_data.get('title', '')}"
        p = self.doc.add_heading(title, level=level)
        p.style = heading_style

        # 添加内容
        content = chapter_data.get('content', '')
        if content:
            # 处理段落
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = self.doc.add_paragraph(para.strip())
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    p.paragraph_format.line_spacing = 1.25
                    p.paragraph_format.first_line_indent = Pt(0)

        # 处理子章节（支持children或subsections键）
        subsections = chapter_data.get('children') or chapter_data.get('subsections', [])
        if subsections:
            for child in subsections:
                self.add_chapter(child, level + 1)

    def add_references(self, references):
        """添加参考文献"""
        # 标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('参考文献')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 文献列表
        for i, ref in enumerate(references, 1):
            p = self.doc.add_paragraph()
            run = p.add_run(f'[{i}] {ref}')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Cm(1)  # 悬挂缩进


def load_thesis_data(project_path):
    """加载论文数据"""
    paper_dir = Path(project_path) / 'paper'

    # 读取章节数据
    chapters_file = paper_dir / 'chapters' / 'chapters.json'
    if not chapters_file.exists():
        print(f"❌ 未找到章节文件: {chapters_file}")
        return None

    with open(chapters_file, 'r', encoding='utf-8') as f:
        chapters_data = json.load(f)
        # 如果JSON包含chapters键，提取数组；否则假设整个JSON就是章节数组
        if isinstance(chapters_data, dict) and 'chapters' in chapters_data:
            chapters = chapters_data['chapters']
        else:
            chapters = chapters_data

    # 读取论文信息
    info_file = paper_dir / 'thesis-info.json'
    if info_file.exists():
        with open(info_file, 'r', encoding='utf-8') as f:
            info = json.load(f)
    else:
        info = {}

    return {
        'chapters': chapters,
        'info': info
    }


def export_to_word(project_path, output_path=None):
    """导出为Word文档"""
    print(f"📦 开始导出论文...")
    print(f"📂 项目路径: {project_path}")

    # 加载数据
    data = load_thesis_data(project_path)
    if not data:
        return False

    # 创建Word文档
    doc = Document()
    formatter = ThesisFormatter(doc)

    # 构建文档
    info = data['info']
    chapters = data['chapters']

    print("✍️  生成封面...")
    formatter.add_cover(info)

    print("✍️  生成声明页...")
    formatter.add_declarations()

    print("✍️  生成目录...")
    formatter.add_toc(chapters)

    # 中文摘要
    if 'abstract_cn' in info:
        print("✍️  生成中文摘要...")
        formatter.add_abstract_cn(
            info['abstract_cn'],
            info.get('keywords_cn', '')
        )

    # 英文摘要
    if 'abstract_en' in info:
        print("✍️  生成英文摘要...")
        formatter.add_abstract_en(
            info['abstract_en'],
            info.get('keywords_en', '')
        )

    # 正文章节
    print("✍️  生成正文章节...")
    for chapter in chapters:
        formatter.add_chapter(chapter)

    # 参考文献
    if 'references' in info:
        print("✍️  生成参考文献...")
        formatter.add_references(info['references'])

    # 保存文档
    if output_path is None:
        output_path = Path(project_path) / 'paper' / '毕业论文.docx'

    doc.save(output_path)

    print(f"\n✅ 论文导出成功！")
    print(f"📄 文件位置: {output_path}")

    return True


def main():
    """命令行入口"""
    if len(sys.argv) > 1:
        project_path = sys.argv[1]
    else:
        project_path = Path.cwd()

    success = export_to_word(project_path)
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()