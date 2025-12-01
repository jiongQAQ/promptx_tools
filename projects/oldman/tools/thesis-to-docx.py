#!/usr/bin/env python3
"""
论文统一导出工具 - 基于样式配置和章节JSON生成Word文档

特性:
- 支持自定义样式配置（docx-styles.json）
- 自动从outline.json读取章节结构
- 从独立的章节文件中加载内容
- 根据docx_type字段应用对应样式
- 支持图片和表格引用

使用方法:
    python3 thesis-to-docx.py [--style STYLE_FILE] [--output OUTPUT_FILE]

示例:
    python3 thesis-to-docx.py
    python3 thesis-to-docx.py --style ../templates/docx-styles-yxnu.json
    python3 thesis-to-docx.py --output 食堂评价系统论文.docx
"""

import json
import os
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


class StyleManager:
    """样式管理器"""

    def __init__(self, style_config_path):
        """初始化样式管理器"""
        with open(style_config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
        self.styles = self.config.get('styles', {})
        self.preset = self.config.get('presets', {}).get('yxnu_thesis', {})

    def get_style(self, style_name):
        """获取样式配置"""
        return self.styles.get(style_name, self.styles.get('default'))

    def apply_style_to_paragraph(self, paragraph, style_name):
        """将样式应用到段落"""
        style_config = self.get_style(style_name)
        if not style_config:
            return

        font_config = style_config.get('font', {})

        # 设置字体
        for run in paragraph.runs:
            # 英文字体
            run.font.name = font_config.get('name', 'Times New Roman')
            # 中文字体
            if 'name_cn' in font_config:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_config['name_cn'])
            # 字号 (half-points)
            if 'size' in font_config:
                run.font.size = Pt(font_config['size'] / 2)
            # 颜色
            if 'color' in font_config:
                run.font.color.rgb = RGBColor.from_string(font_config['color'])
            # 加粗
            if font_config.get('bold'):
                run.font.bold = True
            # 斜体
            if font_config.get('italic'):
                run.font.italic = True

        # 设置对齐
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justified': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        alignment = style_config.get('alignment', 'left')
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # 设置间距
        spacing = style_config.get('spacing', {})
        if 'before' in spacing:
            paragraph.paragraph_format.space_before = Pt(spacing['before'] / 20)
        if 'after' in spacing:
            paragraph.paragraph_format.space_after = Pt(spacing['after'] / 20)
        if 'line' in spacing:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = spacing['line']

        # 设置缩进
        indent = style_config.get('indent', {})
        if 'firstLine' in indent:
            paragraph.paragraph_format.first_line_indent = Pt(indent['firstLine'] / 20)
        if 'left' in indent:
            paragraph.paragraph_format.left_indent = Pt(indent['left'] / 20)
        if 'right' in indent:
            paragraph.paragraph_format.right_indent = Pt(indent['right'] / 20)


class ThesisBuilder:
    """论文构建器"""

    def __init__(self, project_root, style_manager):
        self.project_root = Path(project_root)
        self.style_manager = style_manager
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """设置页面格式"""
        preset = self.style_manager.preset
        if not preset:
            return

        section = self.doc.sections[0]
        page_config = preset.get('page', {})

        # 页边距
        margins = page_config.get('margin', {})
        if 'top' in margins:
            section.top_margin = Pt(margins['top'] / 20)
        if 'bottom' in margins:
            section.bottom_margin = Pt(margins['bottom'] / 20)
        if 'left' in margins:
            section.left_margin = Pt(margins['left'] / 20)
        if 'right' in margins:
            section.right_margin = Pt(margins['right'] / 20)

        # 页眉页脚距离
        if 'header_distance' in page_config:
            section.header_distance = Pt(page_config['header_distance'] / 20)
        if 'footer_distance' in page_config:
            section.footer_distance = Pt(page_config['footer_distance'] / 20)

        # 纸张大小 (A4)
        if page_config.get('size') == 'A4':
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)

    def load_outline(self):
        """加载论文大纲"""
        outline_file = self.project_root / 'paper' / 'outline.json'
        with open(outline_file, 'r', encoding='utf-8') as f:
            outline = json.load(f)
        return outline.get('nodes', [])

    def load_chapter(self, chapter_id):
        """加载章节内容"""
        chapter_file = self.project_root / 'paper' / 'chapters' / f'chapter.{chapter_id}.json'
        if not chapter_file.exists():
            return None

        with open(chapter_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def add_chapter_content(self, chapter_data):
        """添加章节内容"""
        if not chapter_data:
            return

        chapter_id = chapter_data.get('id', '')
        title = chapter_data.get('title', '')
        content = chapter_data.get('content', '')

        # 获取样式类型
        title_style = chapter_data.get('docx_type', 'body_text')
        text_style = chapter_data.get('docx_type_text', 'body_text')

        # 添加标题
        if title:
            # 对于一级章节，添加格式化的标题
            if '.' not in chapter_id and chapter_id not in ['0.1', '0.2']:
                full_title = f"第{chapter_id}章 {title}"
            else:
                full_title = f"{chapter_id} {title}" if chapter_id not in ['0.1', '0.2'] else title

            p = self.doc.add_paragraph(full_title)
            self.style_manager.apply_style_to_paragraph(p, title_style)

        # 添加内容
        if content:
            # 按段落分割
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = self.doc.add_paragraph(para_text)
                    self.style_manager.apply_style_to_paragraph(p, text_style)

    def build_from_outline(self, outline_nodes):
        """根据大纲构建论文"""
        def process_node(node):
            """递归处理节点"""
            node_id = node.get('id', '')

            # 加载并添加章节内容
            chapter_data = self.load_chapter(node_id)
            if chapter_data:
                self.add_chapter_content(chapter_data)
                print(f"  ✅ {node_id} {node.get('title', '')}")
            else:
                print(f"  ⚠️  {node_id} {node.get('title', '')} - 未找到章节文件")

            # 处理子节点
            children = node.get('children', [])
            for child in children:
                process_node(child)

            # 一级章节后添加分页
            if '.' not in node_id and node_id not in ['0.1', '0.2']:
                self.doc.add_page_break()

        # 处理所有节点
        for node in outline_nodes:
            process_node(node)

    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='论文导出工具')
    parser.add_argument('--style', default=None, help='样式配置文件路径')
    parser.add_argument('--output', default=None, help='输出文件路径')
    parser.add_argument('--project', default=None, help='项目根目录')
    args = parser.parse_args()

    # 确定项目根目录
    if args.project:
        project_root = Path(args.project)
    else:
        project_root = Path(__file__).parent.parent

    # 确定样式配置文件
    if args.style:
        style_file = Path(args.style)
    else:
        style_file = project_root / 'templates' / 'docx-styles-yxnu.json'

    if not style_file.exists():
        print(f"❌ 样式配置文件不存在: {style_file}")
        return 1

    # 确定输出文件
    if args.output:
        output_file = Path(args.output)
    else:
        output_file = project_root / 'paper' / '食堂评价系统论文.docx'

    print("📚 论文导出工具")
    print("=" * 60)
    print(f"📂 项目目录: {project_root}")
    print(f"🎨 样式配置: {style_file.name}")
    print(f"📄 输出文件: {output_file}")
    print()

    # 加载样式管理器
    print("⚙️  加载样式配置...")
    style_manager = StyleManager(style_file)
    print(f"  ✅ 已加载 {len(style_manager.styles)} 个样式")
    print()

    # 创建论文构建器
    print("🏗️  构建论文...")
    builder = ThesisBuilder(project_root, style_manager)

    # 加载大纲
    outline_nodes = builder.load_outline()
    print(f"  ✅ 已加载大纲，共 {len(outline_nodes)} 个顶层章节")
    print()

    # 根据大纲构建论文
    print("✍️  生成章节内容...")
    builder.build_from_outline(outline_nodes)
    print()

    # 保存文档
    print("💾 保存文档...")
    output_file.parent.mkdir(parents=True, exist_ok=True)
    builder.save(output_file)

    print()
    print("=" * 60)
    print(f"✅ 论文导出成功！")
    print(f"📄 文件位置: {output_file}")
    print(f"📊 文件大小: {output_file.stat().st_size / 1024:.1f} KB")

    return 0


if __name__ == '__main__':
    sys.exit(main())
