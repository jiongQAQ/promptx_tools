#!/usr/bin/env python3
"""
论文统一导出工具（增强版）- 支持图片和表格

特性:
- 支持自定义样式配置
- 自动从outline.json读取章节结构
- 支持路径变量解析（${er}, ${uml}, ${dfd}, ${flow}, ${tables}）
- 支持图片插入（PNG, SVG）
- 支持表格生成（从JSON）
- 处理items数组

使用方法:
    python3 thesis-to-docx-enhanced.py [--style STYLE_FILE] [--output OUTPUT_FILE]
"""

import json
import os
import sys
import argparse
import tempfile
import subprocess
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Check for SVG conversion tool
SVG_SUPPORT = bool(shutil.which('rsvg-convert'))


class PathResolver:
    """路径变量解析器"""

    def __init__(self, project_root):
        self.project_root = Path(project_root)
        self.path_vars = {
            'er': self.project_root / 'paper' / 'assets' / 'diagrams' / 'er',
            'uml': self.project_root / 'paper' / 'assets' / 'diagrams' / 'uml',
            'dfd': self.project_root / 'paper' / 'assets' / 'diagrams' / 'dfd',
            'flow': self.project_root / 'paper' / 'assets' / 'diagrams' / 'flow',
            'tables': self.project_root / 'paper' / 'assets' / 'tables',
        }

    def resolve(self, path_str):
        """解析路径变量"""
        if not path_str or not isinstance(path_str, str):
            return None

        # 处理 ${var}/file.ext 格式
        if '${' in path_str:
            for var_name, var_path in self.path_vars.items():
                pattern = f'${{{var_name}}}'
                if pattern in path_str:
                    resolved = path_str.replace(pattern, str(var_path))
                    return Path(resolved)

        # 已经是完整路径
        return Path(path_str)


class StyleManager:
    """样式管理器"""

    def __init__(self, style_config_path):
        with open(style_config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
        self.styles = self.config.get('styles', {})
        self.preset = self.config.get('presets', {}).get('yxnu_thesis', {})

    def get_style(self, style_name):
        """获取样式配置"""
        return self.styles.get(style_name, self.styles.get('default'))

    def apply_style_to_paragraph(self, paragraph, style_name):
        """将样式应用到段落"""
        style_config = self.get_style(style_name)
        if not style_config:
            return

        font_config = style_config.get('font', {})

        # 设置字体
        for run in paragraph.runs:
            run.font.name = font_config.get('name', 'Times New Roman')
            if 'name_cn' in font_config:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_config['name_cn'])
            if 'size' in font_config:
                run.font.size = Pt(font_config['size'] / 2)
            if 'color' in font_config:
                run.font.color.rgb = RGBColor.from_string(font_config['color'])
            if font_config.get('bold'):
                run.font.bold = True
            if font_config.get('italic'):
                run.font.italic = True

        # 设置对齐
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justified': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        alignment = style_config.get('alignment', 'left')
        paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # 设置间距
        spacing = style_config.get('spacing', {})
        if 'before' in spacing:
            paragraph.paragraph_format.space_before = Pt(spacing['before'] / 20)
        if 'after' in spacing:
            paragraph.paragraph_format.space_after = Pt(spacing['after'] / 20)
        if 'line' in spacing:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = spacing['line']

        # 设置缩进
        indent = style_config.get('indent', {})
        if 'firstLine' in indent:
            paragraph.paragraph_format.first_line_indent = Pt(indent['firstLine'] / 20)
        if 'left' in indent:
            paragraph.paragraph_format.left_indent = Pt(indent['left'] / 20)
        if 'right' in indent:
            paragraph.paragraph_format.right_indent = Pt(indent['right'] / 20)

        # 设置大纲级别（用于生成目录）
        if 'headingLevel' in style_config:
            level = style_config['headingLevel'] - 1  # Word的级别从0开始
            pPr = paragraph._element.get_or_add_pPr()
            outlineLvl = OxmlElement('w:outlineLvl')
            outlineLvl.set(qn('w:val'), str(level))
            pPr.append(outlineLvl)


class ThesisBuilder:
    """论文构建器"""

    def __init__(self, project_root, style_manager):
        self.project_root = Path(project_root)
        self.style_manager = style_manager
        self.path_resolver = PathResolver(project_root)
        self.doc = Document()
        self._setup_page()

        # 图表计数器 {章节号: {figure: 计数, table: 计数}}
        self.figure_counters = {}
        self.table_counters = {}

    def _setup_page(self):
        """设置页面格式"""
        preset = self.style_manager.preset
        if not preset:
            return

        section = self.doc.sections[0]
        page_config = preset.get('page', {})

        margins = page_config.get('margin', {})
        if 'top' in margins:
            section.top_margin = Pt(margins['top'] / 20)
        if 'bottom' in margins:
            section.bottom_margin = Pt(margins['bottom'] / 20)
        if 'left' in margins:
            section.left_margin = Pt(margins['left'] / 20)
        if 'right' in margins:
            section.right_margin = Pt(margins['right'] / 20)

        if 'header_distance' in page_config:
            section.header_distance = Pt(page_config['header_distance'] / 20)
        if 'footer_distance' in page_config:
            section.footer_distance = Pt(page_config['footer_distance'] / 20)

        if page_config.get('size') == 'A4':
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)

    def get_chapter_number(self, chapter_id):
        """从章节ID提取章节号（用于图表编号）"""
        if not chapter_id:
            return None
        # 提取第一个数字作为章节号，如 "4.2.1" -> "4"
        parts = chapter_id.split('.')
        if parts and parts[0].isdigit():
            return parts[0]
        return None

    def get_title_level_and_style(self, chapter_id):
        """根据章节ID判断标题层级和样式"""
        if not chapter_id:
            return None, None

        # 特殊处理摘要和Abstract
        if chapter_id in ['0.1', '0.2']:
            return 0, 'abstract_title'

        # 统计点号数量来判断层级
        dot_count = chapter_id.count('.')

        if dot_count == 0:
            # 一级标题：1, 2, 3
            return 1, 'chapter_title'
        elif dot_count == 1:
            # 二级标题：1.1, 2.1, 4.1
            return 2, 'section_title'
        elif dot_count == 2:
            # 三级标题：1.1.1, 4.1.2
            return 3, 'subsection_title'
        else:
            # 更深层级，默认使用三级标题样式
            return 3, 'subsection_title'

    def get_next_figure_number(self, chapter_num):
        """获取下一个图编号"""
        if chapter_num not in self.figure_counters:
            self.figure_counters[chapter_num] = 0
        self.figure_counters[chapter_num] += 1
        return f"{chapter_num}-{self.figure_counters[chapter_num]}"

    def get_next_table_number(self, chapter_num):
        """获取下一个表编号"""
        if chapter_num not in self.table_counters:
            self.table_counters[chapter_num] = 0
        self.table_counters[chapter_num] += 1
        return f"{chapter_num}-{self.table_counters[chapter_num]}"

    def load_outline(self):
        """加载论文大纲"""
        outline_file = self.project_root / 'paper' / 'outline.json'
        with open(outline_file, 'r', encoding='utf-8') as f:
            outline = json.load(f)
        return outline.get('outline', [])

    def load_chapter(self, chapter_id):
        """加载章节内容"""
        chapter_file = self.project_root / 'paper' / 'chapters' / f'chapter.{chapter_id}.json'
        if not chapter_file.exists():
            return None

        with open(chapter_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def insert_image(self, image_path, caption=None, width_cm=14):
        """插入图片"""
        if not image_path or not image_path.exists():
            print(f"  ⚠️  图片不存在: {image_path}")
            return

        try:
            # 检查是否是SVG文件
            temp_png_path = None
            if image_path.suffix.lower() == '.svg':
                if not SVG_SUPPORT:
                    print(f"    ⚠️  SVG支持未安装，跳过: {image_path.name}")
                    return

                # 转换SVG到临时PNG文件 (使用rsvg-convert)
                temp_png_fd, temp_png_path = tempfile.mkstemp(suffix='.png')
                os.close(temp_png_fd)

                subprocess.run([
                    'rsvg-convert',
                    '-d', '300',  # DPI
                    '-p', '300',  # DPI
                    '-o', temp_png_path,
                    str(image_path)
                ], check=True, capture_output=True)

                actual_image_path = temp_png_path
                print(f"    🔄 转换SVG: {image_path.name} → PNG")
            else:
                actual_image_path = str(image_path)

            # 插入图片
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(actual_image_path, width=Cm(width_cm))

            # 清理临时文件
            if temp_png_path:
                try:
                    os.unlink(temp_png_path)
                except:
                    pass

            # 添加图题
            if caption:
                p = self.doc.add_paragraph(caption)
                self.style_manager.apply_style_to_paragraph(p, 'figure_caption')

            print(f"    ✅ 插入图片: {image_path.name}")
        except Exception as e:
            print(f"    ❌ 插入图片失败 {image_path.name}: {e}")

    def insert_table(self, table_path, caption=None):
        """插入表格"""
        if not table_path or not table_path.exists():
            print(f"  ⚠️  表格数据不存在: {table_path}")
            return

        try:
            with open(table_path, 'r', encoding='utf-8') as f:
                table_data = json.load(f)

            # 添加表题
            if caption:
                p = self.doc.add_paragraph(caption)
                self.style_manager.apply_style_to_paragraph(p, 'table_caption')

            # 获取列数据
            columns = table_data.get('columns', [])
            if not columns:
                print(f"    ⚠️  表格数据为空")
                return

            # 创建表格
            table = self.doc.add_table(rows=len(columns), cols=len(columns[0]))
            table.style = 'Table Grid'

            # 填充表格数据
            for row_idx, row_data in enumerate(columns):
                row_cells = table.rows[row_idx].cells
                for col_idx, cell_data in enumerate(row_data):
                    cell = row_cells[col_idx]
                    cell.text = str(cell_data)

                    # 设置表头样式（第一行）
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            self.style_manager.apply_style_to_paragraph(paragraph, 'table_header')
                    else:
                        for paragraph in cell.paragraphs:
                            self.style_manager.apply_style_to_paragraph(paragraph, 'table_cell')

            print(f"    ✅ 插入表格: {table_path.name} ({len(columns)}行 × {len(columns[0])}列)")

        except Exception as e:
            print(f"    ❌ 插入表格失败 {table_path.name}: {e}")

    def add_chapter_content(self, chapter_data):
        """添加章节内容"""
        if not chapter_data:
            return

        chapter_id = chapter_data.get('id', '')
        title = chapter_data.get('title', '')
        content = chapter_data.get('content', '')

        # 获取章节号（用于图表编号）
        chapter_num = self.get_chapter_number(chapter_id)

        # 获取样式类型
        title_style = chapter_data.get('docx_type', 'body_text')
        text_style = chapter_data.get('docx_type_text', 'body_text')

        # 添加标题
        if title:
            if '.' not in chapter_id and chapter_id not in ['0.1', '0.2']:
                full_title = f"第{chapter_id}章 {title}"
            else:
                full_title = f"{chapter_id} {title}" if chapter_id not in ['0.1', '0.2'] else title

            p = self.doc.add_paragraph(full_title)
            self.style_manager.apply_style_to_paragraph(p, title_style)

        # 添加内容
        if content:
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = self.doc.add_paragraph(para_text)
                    self.style_manager.apply_style_to_paragraph(p, text_style)

        # 处理关键词（摘要和Abstract专用）
        keywords = chapter_data.get('keywords', '')
        if keywords:
            # 中文摘要：关键词：xxx；xxx；xxx
            # 英文摘要：Key words: xxx; xxx; xxx
            if chapter_id == '0.1':
                keyword_text = f"关键词：{keywords}"
            elif chapter_id == '0.2':
                keyword_text = f"Key words: {keywords}"
            else:
                keyword_text = f"关键词：{keywords}"

            p = self.doc.add_paragraph(keyword_text)
            self.style_manager.apply_style_to_paragraph(p, 'keywords')

        # 处理items数组（包含子项的章节）
        items = chapter_data.get('items', [])
        for item in items:
            item_title = item.get('title', '')
            item_text = item.get('text', '')

            # 添加子项标题
            if item_title:
                p = self.doc.add_paragraph(item_title)
                self.style_manager.apply_style_to_paragraph(p, 'subsection_title')

            # 添加子项内容
            if item_text:
                p = self.doc.add_paragraph(item_text)
                self.style_manager.apply_style_to_paragraph(p, text_style)

            # 处理子项图片
            if 'imagePath' in item:
                image_path = self.path_resolver.resolve(item['imagePath'])
                if image_path and image_path.exists():
                    # 生成图编号和标题
                    if chapter_num:
                        fig_num = self.get_next_figure_number(chapter_num)
                        caption = f"图 {fig_num}  {item_title}"  # 编号后空两个字
                    else:
                        caption = f"图  {item_title}"
                    self.insert_image(image_path, caption=caption)

            # 处理子项表格
            if 'tablePath' in item:
                table_path = self.path_resolver.resolve(item['tablePath'])
                if table_path and table_path.exists():
                    # 生成表编号和标题
                    if chapter_num:
                        tab_num = self.get_next_table_number(chapter_num)
                        caption = f"表 {tab_num}  {item_title}"  # 编号后空两个字
                    else:
                        caption = f"表  {item_title}"
                    self.insert_table(table_path, caption=caption)

        # 处理章节级别的图片
        if 'imagePath' in chapter_data:
            image_path = self.path_resolver.resolve(chapter_data['imagePath'])
            if image_path and image_path.exists():
                # 生成图编号和标题
                if chapter_num:
                    fig_num = self.get_next_figure_number(chapter_num)
                    caption = f"图 {fig_num}  {title}"
                else:
                    caption = title if title else None
                self.insert_image(image_path, caption=caption)

        # 处理章节级别的表格
        if 'tablePath' in chapter_data:
            table_path = self.path_resolver.resolve(chapter_data['tablePath'])
            if table_path and table_path.exists():
                # 生成表编号和标题
                if chapter_num:
                    tab_num = self.get_next_table_number(chapter_num)
                    caption = f"表 {tab_num}  {title}"
                else:
                    caption = title if title else None
                self.insert_table(table_path, caption=caption)

    def build_from_outline(self, outline_nodes):
        """根据大纲构建论文"""
        # 构建ID到节点的映射
        node_map = {node['id']: node for node in outline_nodes}

        def process_node(node):
            node_id = node.get('id', '')
            node_title = node.get('title', '')

            # 加载并添加章节内容
            chapter_data = self.load_chapter(node_id)
            if chapter_data:
                self.add_chapter_content(chapter_data)
                print(f"  ✅ {node_id} {node_title}")
            else:
                # 章节文件不存在时，根据层级添加对应的标题
                level, style_name = self.get_title_level_and_style(node_id)

                if level is not None and style_name and node_id not in ['0.1', '0.2']:
                    # 格式化标题：添加章节编号
                    if level == 1:
                        # 一级标题已经包含"第X章"，直接使用
                        formatted_title = f"第{node_id}章 {node_title}"
                    elif level == 2:
                        # 二级标题：5.1 标题
                        formatted_title = f"{node_id} {node_title}"
                    elif level == 3:
                        # 三级标题：5.1.1 标题
                        formatted_title = f"{node_id} {node_title}"
                    else:
                        formatted_title = f"第{node_id}章 {node_title}"

                    # 添加标题段落
                    p = self.doc.add_paragraph(formatted_title)
                    self.style_manager.apply_style_to_paragraph(p, style_name)
                    print(f"  ✅ {formatted_title} (使用大纲标题)")
                else:
                    print(f"  ⚠️  {node_id} {node_title} - 未找到章节文件")

            # 处理子节点（children是ID字符串数组）
            children = node.get('children', [])
            for child_id in children:
                if child_id in node_map:
                    process_node(node_map[child_id])

            # 一级章节后添加分页
            if '.' not in node_id and node_id not in ['0.1', '0.2']:
                self.doc.add_page_break()

        # 只处理顶层节点（避免重复）
        top_nodes = [n for n in outline_nodes if n.get('parent') is None]
        for node in top_nodes:
            process_node(node)

    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='论文导出工具（增强版）')
    parser.add_argument('--style', default=None, help='样式配置文件路径')
    parser.add_argument('--output', default=None, help='输出文件路径')
    parser.add_argument('--project', default=None, help='项目根目录')
    args = parser.parse_args()

    # 确定项目根目录
    if args.project:
        project_root = Path(args.project)
    else:
        project_root = Path(__file__).parent.parent

    # 确定样式配置文件
    if args.style:
        style_file = Path(args.style)
    else:
        style_file = project_root / 'templates' / 'docx-styles-yxnu.json'

    if not style_file.exists():
        print(f"❌ 样式配置文件不存在: {style_file}")
        return 1

    # 确定输出文件
    if args.output:
        output_file = Path(args.output)
    else:
        output_file = project_root / 'paper' / '食堂评价系统论文-完整版.docx'

    print("📚 论文导出工具（增强版）")
    print("=" * 60)
    print(f"📂 项目目录: {project_root}")
    print(f"🎨 样式配置: {style_file.name}")
    print(f"📄 输出文件: {output_file}")
    print()

    # 加载样式管理器
    print("⚙️  加载样式配置...")
    style_manager = StyleManager(style_file)
    print(f"  ✅ 已加载 {len(style_manager.styles)} 个样式")
    print()

    # 创建论文构建器
    print("🏗️  构建论文...")
    builder = ThesisBuilder(project_root, style_manager)

    # 加载大纲
    outline_nodes = builder.load_outline()
    print(f"  ✅ 已加载大纲，共 {len(outline_nodes)} 个顶层章节")
    print()

    # 根据大纲构建论文
    print("✍️  生成章节内容（包含图片和表格）...")
    builder.build_from_outline(outline_nodes)
    print()

    # 保存文档
    print("💾 保存文档...")
    output_file.parent.mkdir(parents=True, exist_ok=True)
    builder.save(output_file)

    print()
    print("=" * 60)
    print(f"✅ 论文导出成功！")
    print(f"📄 文件位置: {output_file}")
    print(f"📊 文件大小: {output_file.stat().st_size / 1024:.1f} KB")

    return 0


if __name__ == '__main__':
    sys.exit(main())
